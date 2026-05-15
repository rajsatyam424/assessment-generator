"""
DOCX builder — produces a formatted Word document from assessment data.
"""

import os
import re
import logging
from pathlib import Path
from typing import Any

from docx import Document
from docx.shared import Pt, Inches

from . import config

logger = logging.getLogger(__name__)


def build_docx(data: dict[str, Any]) -> str:
    """
    Build a formatted .docx file from structured assessment data.

    Args:
        data: Parsed assessment dict (from parser.parse_response).

    Returns:
        Absolute path to the generated .docx file.
    """
    doc = Document()
    section = doc.sections[0]
    margin = Inches(config.MARGIN_INCHES)
    for attr in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, attr, margin)

    def _heading(text: str, size: int):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.bold = True
        r.font.name = config.FONT_NAME
        r.font.size = Pt(size)
        return p

    def _para(text: str, bold: bool = False, indent: bool = False):
        p = doc.add_paragraph()
        if indent:
            p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(text)
        r.bold = bold
        r.font.name = config.FONT_NAME
        r.font.size = Pt(config.FONT_SIZE)
        return p

    def _bullet(text: str):
        p = doc.add_paragraph(style="List Bullet")
        p.clear()
        r = p.add_run(text)
        r.font.name = config.FONT_NAME
        r.font.size = Pt(config.FONT_SIZE)
        return p

    # --- Title ---
    _heading(data.get("assessment_name", "Untitled Assessment"), 18)

    # --- About ---
    _heading("About this assessment", 14)
    for p_text in data.get("about_assessment", []):
        _para(p_text)

    # --- Audience ---
    _heading("Who this is for", 14)
    _para(data.get("who_this_is_for", ""))

    # --- Learning Outcomes ---
    _heading("Learning Outcomes", 14)
    _para("By the end of this assessment, you should be able to:")
    for lo in data.get("learning_outcomes", []):
        _bullet(lo)

    # --- Topics ---
    _heading("Topics covered", 14)
    for t in data.get("topics", []):
        _bullet(t)

    # --- Metadata ---
    _para(f"Total questions: {len(data.get('questions', []))}", bold=True)
    _para("Format: Multiple choice", bold=True)
    _para(f"Estimated time: {data.get('estimated_time', '30 minutes')}", bold=True)

    # --- Questions ---
    for q in data.get("questions", []):
        n = q.get("number", "?")
        stem = q.get("stem", "")
        topic = q.get("topic", "")
        _para(f"Q{n}. {topic} {stem}", bold=True)
        _para(f"A) {q.get('a', '')}", indent=True)
        _para(f"B) {q.get('b', '')}", indent=True)
        _para(f"C) {q.get('c', '')}", indent=True)
        _para(f"D) {q.get('d', '')}", indent=True)
        _para(f"Answer: {q.get('answer', '?')}")
        _para(f"Explanation: {q.get('explanation', '')}")

    # --- Save ---
    config.OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    safe_name = re.sub(r"[^a-zA-Z0-9\s]", "", data.get("assessment_name", "assessment")).strip().replace(" ", "_")
    filename = f"{safe_name}_{os.urandom(4).hex()}.docx"
    output_path = str(config.OUTPUT_DIR / filename)
    doc.save(output_path)

    logger.info("DOCX saved: %s", output_path)
    return output_path
